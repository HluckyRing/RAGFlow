import os
import io
import csv
from pypdf import PdfReader
from src.config import logger

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False


def load_pdf(uploaded_file):
    reader = PdfReader(uploaded_file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text
    return text


def load_txt(uploaded_file):
    content = uploaded_file.read()
    for encoding in ('utf-8', 'gbk', 'gb2312', 'latin-1'):
        try:
            return content.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return content.decode('utf-8', errors='replace')


def load_docx(uploaded_file):
    if not HAS_DOCX:
        raise ImportError("请安装 python-docx: pip install python-docx")
    doc = Document(uploaded_file)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs)


def load_xlsx(uploaded_file):
    if not HAS_OPENPYXL:
        raise ImportError("请安装 openpyxl: pip install openpyxl")
    wb = openpyxl.load_workbook(uploaded_file, read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"=== {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
            if row_text.strip():
                parts.append(row_text)
    wb.close()
    return "\n".join(parts)


def load_csv(uploaded_file):
    content = uploaded_file.read().decode('utf-8')
    reader = csv.reader(io.StringIO(content))
    rows = [" | ".join(row) for row in reader if any(cell.strip() for cell in row)]
    return "\n".join(rows)


LOADERS = {
    '.pdf': load_pdf,
    '.txt': load_txt,
    '.md': load_txt,
    '.docx': load_docx,
    '.xlsx': load_xlsx,
    '.csv': load_csv,
}

SUPPORTED_EXTENSIONS = list(LOADERS.keys())
SUPPORTED_TYPES = [ext.lstrip('.') for ext in SUPPORTED_EXTENSIONS]


def load_file(uploaded_file):
    filename = getattr(uploaded_file, 'name', None) or getattr(uploaded_file, 'filename', '')
    ext = os.path.splitext(filename)[1].lower()
    loader = LOADERS.get(ext)
    if loader is None:
        raise ValueError(f"不支持的文件格式: {ext or '未知'}，支持: {', '.join(SUPPORTED_EXTENSIONS)}")
    logger.info("加载文件: %s (%s)", filename, ext)
    return loader(uploaded_file)
